#!/usr/bin/env python3
"""Generate an editable Word report from the Redis RAG markdown report."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_MD = ROOT / "report" / "redis_rag_report.md"
OUT_DOCX = ROOT / "report" / "redis_rag_report.docx"
ASSETS = ROOT / "slides" / "assets"

FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"
ACCENT = RGBColor(217, 45, 32)
INK = RGBColor(23, 33, 43)
MUTED = RGBColor(100, 116, 139)
TABLE_FILL = "FDE8E5"


def set_run_font(run, size: float | None = None, bold: bool | None = None, color=None, name: str = FONT_CN) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(clean_inline(text))
    set_run_font(run, size=9.5, bold=bold, color=INK)


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    return text.strip()


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_CN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, INK),
        ("Heading 3", 11.5, INK),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT_CN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(5)

    code = doc.styles.add_style("CodeBlock", 1)
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    code.font.size = Pt(8.7)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("基于 RAG 的 Redis 垂直领域深度问答系统构建")
    set_run_font(run, size=22, bold=True, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("课程作业报告 | 私有知识库、混合检索、轻量重排序与三维量化评估")
    set_run_font(run, size=11.5, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("生成日期：2026-06-17    默认 LLM：deepseek-v4-pro    可编辑版本：DOCX")
    set_run_font(run, size=10, color=MUTED)

    doc.add_paragraph()


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            if r_idx == 0:
                shade_cell(cell, TABLE_FILL)
            set_cell_text(cell, value, bold=r_idx == 0)
    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [cell.strip() for cell in raw.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_markdown_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                if code_lines:
                    paragraph = doc.add_paragraph(style="CodeBlock")
                    run = paragraph.add_run("\n".join(code_lines))
                    set_run_font(run, size=8.7, name="Courier New")
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            # Cover page already carries the title.
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(clean_inline(stripped[3:]), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(clean_inline(stripped[4:]), level=2)
            i += 1
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(clean_inline(bullet.group(1)))
            set_run_font(run, size=10.2, color=INK)
        elif numbered:
            paragraph = doc.add_paragraph(style="List Number")
            run = paragraph.add_run(clean_inline(numbered.group(1)))
            set_run_font(run, size=10.2, color=INK)
        else:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(clean_inline(stripped))
            set_run_font(run, size=10.5, color=INK)
        i += 1


def add_figures(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("附录：真实运行结果与展示图片", level=1)
    figures = [
        ("系统架构图", ASSETS / "architecture_flow.png"),
        ("真实索引构建运行结果", ASSETS / "build_index_result.png"),
        ("真实一键推理运行结果", ASSETS / "inference_result.png"),
        ("三维量化评估结果与真实运行输出", ASSETS / "evaluation_results.png"),
        ("检索策略消融对比结果", ASSETS / "retrieval_comparison.png"),
        ("检索命中文档与二阶段重排序分数", ASSETS / "retrieval_scores.png"),
    ]
    for caption, path in figures:
        if not path.exists():
            continue
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(caption)
        set_run_font(run, size=11, bold=True, color=ACCENT)
        doc.add_picture(str(path), width=Inches(6.25))
        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = note.add_run(f"图：{caption}")
        set_run_font(run, size=9, color=MUTED)


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Redis-RAG 课程作业报告")
    set_run_font(run, size=8.5, color=MUTED)


def main() -> None:
    doc = Document()
    apply_styles(doc)
    add_footer(doc)
    add_cover(doc)
    add_markdown_body(doc, REPORT_MD.read_text(encoding="utf-8"))
    add_figures(doc)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
